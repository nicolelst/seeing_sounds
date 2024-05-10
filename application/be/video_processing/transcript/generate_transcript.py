from typing import List, Tuple
from docx import Document
from docx.shared import Inches, RGBColor, Pt

from utils.speaker_info import Speaker_Info
from utils.video_name_constants import get_output_transcript_filename, get_speaker_thumbnail_filepaths, get_whisper_transcript_filepaths
from video_processing.transcript.transcript_iterator import Transcript_Iterator


def generate_transcript(request_id: str, speaker_info: List[Speaker_Info]):
    def add_paragraph(text: str, bold: bool = False, italic: bool = False, underline: bool = False, rgb: Tuple[int, int, int] = None, fontsize:int = 12, indent:float = 0):
        para = document.add_paragraph()
        run = para.add_run(text)

        run.bold = bold
        run.italic = italic
        run.underline = underline

        para.paragraph_format.left_indent = Inches(indent)

        run.font.size = Pt(fontsize)
        
        if rgb != None: 
            run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
    
    def add_speaker_header(idx: int):
        add_paragraph(speaker_info[idx].name.upper(), rgb=speaker_info[idx].colour.as_rgb_tuple(), bold=True, fontsize=14)
    
    def add_section_header(text: str):
        add_paragraph(text, bold=True, underline=True, fontsize=18)
    
    document = Document()

    add_paragraph('Seeing Sounds Transcript', bold=True, fontsize=26) 
    add_paragraph("This transcript was automatically genarated by Seeing Sounds.", italic=True)
    add_paragraph(f"Request ID:\t{request_id}", italic=True, rgb=(136, 136, 136))

    add_section_header('Speakers')
    speaker_images = get_speaker_thumbnail_filepaths(request_id)
    for i in range(len(speaker_info)):
        add_speaker_header(i)
        document.add_picture(speaker_images[i], width=Inches(1.25)) 

    document.add_page_break()

    add_section_header('Transcript')
    transcripts = [Transcript_Iterator(ts) for ts in get_whisper_transcript_filepaths(request_id)] 
    next_segments = [next(ts) for ts in transcripts]
    while not all([n == None for n in next_segments]):
        # print(next_segments)
        # start, end, text
        nextStart = min([seg["start"] for seg in next_segments if seg != None])
        for i in range(len(transcripts)):
            if next_segments[i] != None and next_segments[i]["start"] == nextStart:
                seg = next_segments[i]
                # add segment to doc
                add_speaker_header(i)
                add_paragraph(seg["text"], indent=0.5)
                add_paragraph(f'({seg["start"]:.2f}s - {seg["end"]:.2f}s)', indent=0.5, italic=True, rgb=(136, 136, 136))
                next_segments[i] = next(transcripts[i])
    
    output_filepath = get_output_transcript_filename(request_id)
    document.save(output_filepath)
    return output_filepath
